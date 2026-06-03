import os
import io
import re
import json
import random
from datetime import datetime, timedelta
from dotenv import load_dotenv

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash

# Document Parsers
import fitz  # PyMuPDF
import docx  # python-docx
from PIL import Image

# OCR
try:
    import pytesseract
except ImportError:
    pytesseract = None

# ReportLab for PDF Exports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# AI Integration
import groq

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'intellicard-ai-super-secret-key-1337')
db_url = os.getenv('DATABASE_URL', 'sqlite:///database.db')
if db_url.startswith("postgres://"):
    db_url = db_url.replace("postgres://", "postgresql://", 1)
app.config['SQLALCHEMY_DATABASE_URI'] = db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Setup Auth
login_manager = LoginManager()
login_manager.login_app = app
login_manager.login_view = 'login'
login_manager.init_app(app)

# ──────────────────────────────────────────────────────────
# DATABASE MODELS
# ──────────────────────────────────────────────────────────

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False, index=True)
    password_hash = db.Column(db.String(256), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    decks = db.relationship('Deck', backref='user', lazy=True, cascade="all, delete-orphan")
    quiz_sessions = db.relationship('QuizHistory', backref='user', lazy=True, cascade="all, delete-orphan")

class Deck(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    name = db.Column(db.String(150), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    cards = db.relationship('Card', backref='deck', lazy=True, cascade="all, delete-orphan")

class Card(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    deck_id = db.Column(db.Integer, db.ForeignKey('deck.id'), nullable=False)
    topic = db.Column(db.String(100), default='General')
    question = db.Column(db.Text, nullable=False)
    answer = db.Column(db.Text, nullable=False)
    hint = db.Column(db.Text, nullable=True)
    example = db.Column(db.Text, nullable=True)
    difficulty = db.Column(db.String(20), default='Medium') # 'Easy', 'Medium', 'Hard'
    known_status = db.Column(db.String(20), default='unseen') # 'unseen', 'known', 'review'
    last_reviewed = db.Column(db.DateTime, default=datetime.utcnow)

class QuizHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    deck_id = db.Column(db.Integer, db.ForeignKey('deck.id'), nullable=False)
    score = db.Column(db.Integer, nullable=False)
    total_questions = db.Column(db.Integer, nullable=False)
    mode = db.Column(db.String(20), nullable=False) # 'mcq', 'tf', 'fitb'
    date_taken = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ──────────────────────────────────────────────────────────
# PARSING ENGINE & 4-LEVEL SEGMENTATION
# ──────────────────────────────────────────────────────────

def parse_pdf(stream):
    """
    Parses a PDF stream page-by-page. Returns sections split by detected font changes
    (L1), bold text (L2), or basic page chunks (L4) as text blocks.
    """
    doc = fitz.open(stream=stream, filetype="pdf")
    sections = []
    current_section = {"title": "Introduction", "content": []}
    
    # Track font sizes to find heading threshold
    font_sizes = []
    
    # Phase 1: Collect text details
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" in b:
                for l in b["lines"]:
                    for s in l["spans"]:
                        font_sizes.append(s["size"])

    avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 10
    # Threshold for heading: font size is significantly larger than body text
    heading_threshold = avg_size + 3.0

    # Phase 2: Structural segmentation
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" in b:
                block_text = ""
                is_heading = False
                font_name = ""
                
                # Check if this block functions as a heading
                for l in b["lines"]:
                    for s in l["spans"]:
                        block_text += s["text"] + " "
                        # L1: Font Size check
                        if s["size"] >= heading_threshold:
                            is_heading = True
                        # L2: Bold font flags check (flags bit 4 indicates bold usually, or check font name)
                        if "bold" in s["font"].lower() or "black" in s["font"].lower():
                            # Only treat as header if it's a short line
                            if len(s["text"].strip()) < 80:
                                is_heading = True
                
                block_text = block_text.strip()
                if not block_text:
                    continue

                if is_heading and len(block_text) < 100:
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"title": block_text, "content": []}
                else:
                    current_section["content"].append(block_text)

    if current_section["content"]:
        sections.append(current_section)

    # Flatten list of section dicts into string segments
    flattened_segments = []
    for s in sections:
        header = f"Topic: {s['title']}\n"
        body = "\n".join(s["content"])
        if len(body.strip()) > 50:
            flattened_segments.append(header + body)

    # L4 Fallback: If no headings were detected, split by page
    if not flattened_segments:
        for page in doc:
            text = page.get_text().strip()
            if text:
                flattened_segments.append(text)

    # L5 Fallback: If still no text was extracted, perform OCR page-by-page
    if not flattened_segments and pytesseract:
        print("[*] PDF text extraction yielded no content. Falling back to OCR page-by-page...")
        for page_num, page in enumerate(doc):
            try:
                # Render page to an image pixmap
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # Perform OCR
                ocr_text = pytesseract.image_to_string(img).strip()
                if ocr_text:
                    print(f"[*] Page {page_num + 1} OCR extracted {len(ocr_text)} characters.")
                    flattened_segments.append(ocr_text)
                else:
                    print(f"[*] Page {page_num + 1} OCR returned no text.")
            except Exception as ocr_err:
                print(f"[!] OCR failed on page {page_num + 1}: {ocr_err}")

    return flattened_segments

def parse_docx(stream):
    """
    Parses a DOCX stream. Uses paragraph styles (Heading 1/2/3) or bold tags
    to chunk topics (L1/L2), falling back to character limits if uniform.
    """
    doc = docx.Document(io.BytesIO(stream))
    segments = []
    current_segment = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # L1: Heading styles
        is_heading = p.style.name.startswith('Heading')
        
        # L2: First run bold check
        if not is_heading and p.runs and p.runs[0].bold and len(text) < 100:
            is_heading = True
            
        if is_heading:
            if current_segment:
                segments.append("\n".join(current_segment))
            current_segment = [f"Topic: {text}"]
        else:
            current_segment.append(text)
            
    if current_segment:
        segments.append("\n".join(current_segment))
        
    # Fallback to length chunking if uniform document
    if not segments:
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        segments = chunk_by_length(full_text)
        
    return segments

def parse_image(stream):
    """
    Tries to perform OCR on image. Fallback to warning if tesseract is missing.
    """
    if not pytesseract:
        return ["Warning: pytesseract is not installed on this system. Image upload degraded."]
    try:
        img = Image.open(io.BytesIO(stream))
        ocr_text = pytesseract.image_to_string(img)
        if ocr_text.strip():
            return chunk_by_length(ocr_text)
        return []
    except Exception as e:
        print(f"OCR Error: {e}")
        return []

def chunk_by_length(text, size=1500):
    """
    Basic fallback chunker (L4) that breaks text into page-equivalent sizes.
    """
    paragraphs = text.split('\n')
    chunks = []
    current = []
    length = 0
    for p in paragraphs:
        current.append(p)
        length += len(p)
        if length >= size:
            chunks.append("\n".join(current))
            current = []
            length = 0
    if current:
        chunks.append("\n".join(current))
    return chunks

# ──────────────────────────────────────────────────────────
# FLASHCARD GENERATOR ENGINE (GROQ VS OFFLINE RULE-BASED)
# ──────────────────────────────────────────────────────────

def generate_cards_offline(text):
    """
    Highly robust rule-based parser that scans a text segment for terms and definitions,
    definitional verbs, and key facts to generate educational flashcards locally.
    """
    cards = []
    # Identify title or general theme if marked
    topic_match = re.search(r"Topic:\s*(.*)", text)
    topic = topic_match.group(1).strip() if topic_match else "General"
    
    # Remove topic header line from processing
    clean_text = re.sub(r"Topic:\s*.*\n?", "", text)
    
    # Rule 1: Find term/definition patterns. E.g., Term: Definition or Term - Definition
    # Avoid picking sentences that are too long for a term
    definitions = re.findall(r"(?:^|\n)([\w\s\-\(\)\/\',]{2,40})(?:\s*:\s*|\s+—\s+|\s+-\s+)([^\n.]{15,300}\.?)", clean_text)
    for term, desc in definitions:
        term = term.strip()
        desc = desc.strip()
        if term and desc and len(term) > 2:
            cards.append({
                "topic": topic,
                "question": f"What is the definition of '{term}'?",
                "answer": desc,
                "hint": f"Recall the concept of {term}.",
                "example": None,
                "difficulty": "Easy" if len(desc) < 100 else "Medium"
            })

    # Rule 2: Sentential definitions. E.g. "A neuron is defined as..." or "An API refers to..."
    # Matches words followed by definitional connectors
    connectors = [
        r"is defined as",
        r"refers to",
        r"is a type of",
        r"is characterized by",
        r"consists of",
        r"is the process of"
    ]
    
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', clean_text)
    for sent in sentences:
        sent = sent.strip()
        for conn in connectors:
            if conn in sent:
                parts = sent.split(conn, 1)
                term = parts[0].strip()
                desc = parts[1].strip()
                # Clean up leading/trailing quotes and capitalized starting words
                term = re.sub(r'^(A|An|The)\s+', '', term, flags=re.IGNORECASE)
                if len(term) < 50 and len(desc) > 10:
                    cards.append({
                        "topic": topic,
                        "question": f"Explain what is meant by '{term}'.",
                        "answer": f"{term} {conn} {desc}",
                        "hint": f"Think about {term}.",
                        "example": None,
                        "difficulty": "Medium"
                    })
                    break # Matches only one connector per sentence

    # Rule 3: Key bullets fallback. If no structure is matched, create summarizing questions
    if not cards:
        lines = [line.strip() for line in clean_text.split('\n') if len(line.strip()) > 60]
        for line in lines[:4]: # Limit to 4 cards per segment
            # Generate a question summarizing the paragraph
            first_words = " ".join(line.split()[:4])
            cards.append({
                "topic": topic,
                "question": f"Explain the context regarding: '{first_words}...'",
                "answer": line,
                "hint": "Refer back to the uploaded document text.",
                "example": None,
                "difficulty": "Hard"
            })
            
    return cards

def generate_cards_ai(text, deck_name=None):
    """
    Generates high fidelity flashcards using Groq Llama-3 API.
    If no key is configured, falls back to the offline rule-based extractor.
    """
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        print("No GROQ_API_KEY found. Falling back to local offline extractor.")
        return generate_cards_offline(text)
        
    try:
        client = groq.Groq(api_key=api_key)
        
        prompt = (
            "You are an expert educational tutor. Analyze the following document text and "
            "generate a list of high-quality flashcards to study the core concepts. "
            "For each card, extract or write:\n"
            "1. 'topic': The specific topic or sub-concept (1-3 words).\n"
            "2. 'question': A clear, direct study question.\n"
            "3. 'answer': A concise, accurate, and comprehensive definition or explanation.\n"
            "4. 'hint': An optional, helpful hint (1 sentence).\n"
            "5. 'example': An optional illustrative example, scenario, or code block.\n"
            "6. 'difficulty': A classification ('Easy', 'Medium', 'Hard') based on complexity.\n\n"
            "Respond ONLY with a valid JSON array of objects, where each object has these exact keys: "
            "\"topic\", \"question\", \"answer\", \"hint\", \"example\", \"difficulty\". Do not wrap your response in markdown formatting or add any prose introductory text."
        )

        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Document text to process:\n\n{text}"}
            ],
            model="llama3-8b-8192",
            temperature=0.2,
            response_format={"type": "json_object"}
        )

        raw_content = response.choices[0].message.content.strip()
        
        # Parse JSON output. It might be returned inside a parent key or directly as an array.
        data = json.loads(raw_content)
        if isinstance(data, dict):
            # If model returned something like {"flashcards": [...]}
            for key in data:
                if isinstance(data[key], list):
                    return data[key]
        if isinstance(data, list):
            return data
            
        return generate_cards_offline(text) # Fallback if structure is unexpected
    except Exception as e:
        print(f"AI Generation Error: {e}. Falling back to offline extraction.")
        return generate_cards_offline(text)

# ──────────────────────────────────────────────────────────
# ROUTE DIRECTORIES
# ──────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        
        if not email or not password:
            flash('All fields are required.', 'error')
            return redirect(url_for('signup'))
            
        existing = User.query.filter_by(email=email).first()
        if existing:
            flash('Email already registered.', 'error')
            return redirect(url_for('signup'))
            
        hashed = generate_password_hash(password)
        user = User(email=email, password_hash=hashed)
        db.session.add(user)
        db.session.commit()
        
        login_user(user)
        flash('Account created successfully!', 'success')
        return redirect(url_for('dashboard'))
        
    return render_template('signup.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        
        user = User.query.filter_by(email=email).first()
        if not user or not check_password_hash(user.password_hash, password):
            flash('Invalid email or password.', 'error')
            return redirect(url_for('login'))
            
        login_user(user)
        flash('Successfully logged in!', 'success')
        return redirect(url_for('dashboard'))
        
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully.', 'info')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    decks = Deck.query.filter_by(user_id=current_user.id).all()
    
    # Attach card counts to decks object dynamically
    for d in decks:
        d.cards_count = Card.query.filter_by(deck_id=d.id).count()

    total_cards = sum(d.cards_count for d in decks)
    
    # Calculate stats
    known_cards = 0
    review_cards = 0
    
    for d in decks:
        known_cards += Card.query.filter_by(deck_id=d.id, known_status='known').count()
        review_cards += Card.query.filter_by(deck_id=d.id, known_status='review').count()

    known_pct = Math_pct(known_cards, total_cards)
    review_pct = Math_pct(review_cards, total_cards)

    # 1. Weekly activity logs (cards studied / quiz taken history)
    # Formulate mock or log-based chart data
    today = datetime.utcnow().date()
    labels = []
    counts = []
    
    # Query quiz history logs over last 7 days
    for i in range(6, -1, -1):
        day = today - timedelta(days=i)
        labels.append(day.strftime('%a'))
        
        # Count quizzes completed on this day
        q_count = QuizHistory.query.filter(
            QuizHistory.user_id == current_user.id,
            db.func.date(QuizHistory.date_taken) == day
        ).count()
        
        # Mix in cards updated on this day
        c_count = Card.query.join(Deck).filter(
            Deck.user_id == current_user.id,
            db.func.date(Card.last_reviewed) == day,
            Card.known_status != 'unseen'
        ).count()
        
        counts.append(q_count * 10 + c_count) # Weight quiz completions more

    # 2. Mastery by Topic Analysis
    # Get card groupings by topic
    topic_data = db.session.query(
        Card.topic,
        db.func.count(Card.id).label('total'),
        db.func.sum(db.case((Card.known_status == 'known', 1), else_=0)).label('known')
    ).join(Deck).filter(Deck.user_id == current_user.id).group_by(Card.topic).all()
    
    topic_mastery = []
    weak_topics = []
    colors_classes = ['bar-cyan', 'bar-purple', 'bar-green', 'bar-amber', 'bar-red']
    
    for idx, row in enumerate(topic_data):
        topic = row[0] or 'General'
        total = row[1]
        known = row[2] or 0
        pct = Math_pct(known, total)
        color = colors_classes[idx % len(colors_classes)]
        
        topic_mastery.append({
            "topic": topic,
            "total_count": total,
            "known_count": known,
            "pct": pct,
            "color_class": color
        })
        
        # Diagnose weak topics (mastery < 60%)
        if pct < 60:
            status_text = "Critical" if pct < 30 else "Warning"
            status_class = "critical" if pct < 30 else "warning"
            weak_topics.append({
                "topic": topic,
                "pct": pct,
                "known_count": known,
                "total_count": total,
                "status_text": status_text,
                "status_class": status_class
            })
            
    stats = {
        "total_cards": total_cards,
        "known_cards": known_cards,
        "review_cards": review_cards,
        "known_pct": known_pct,
        "review_pct": review_pct,
        "activity_data": {
            "labels": labels,
            "counts": counts
        },
        "topic_mastery": topic_mastery[:5], # top 5
        "weak_topics": weak_topics[:4] # top 4 weak
    }

    return render_template('dashboard.html', decks=decks, stats=stats)

def Math_pct(val, total):
    if not total:
        return 0
    return round((val / total) * 100)

@app.route('/flashcards')
@login_required
def flashcards():
    deck_id = request.args.get('deck')
    if not deck_id:
        # Fall back to first available deck
        deck = Deck.query.filter_by(user_id=current_user.id).first()
        if not deck:
            flash('You do not have any study decks yet. Upload a file below!', 'info')
            return redirect(url_for('index'))
        deck_id = deck.id
        
    deck = Deck.query.filter_by(id=deck_id, user_id=current_user.id).first_or_404()
    cards = Card.query.filter_by(deck_id=deck.id).all()
    
    # Format cards to simple list of dicts for JSON
    cards_data = []
    for c in cards:
        cards_data.append({
            "id": c.id,
            "topic": c.topic,
            "question": c.question,
            "answer": c.answer,
            "hint": c.hint,
            "example": c.example,
            "difficulty": c.difficulty,
            "status": c.known_status
        })
        
    return render_template('flashcard.html', deck_name=deck.name, deck_id=deck.id, cards=cards_data)

@app.route('/quiz')
@login_required
def quiz():
    selected_deck_id = request.args.get('deck', type=int)
    decks = Deck.query.filter_by(user_id=current_user.id).all()
    for d in decks:
        d.cards_count = Card.query.filter_by(deck_id=d.id).count()
        
    # Filter out empty decks
    active_decks = [d for d in decks if d.cards_count > 0]
    
    return render_template('quiz.html', decks=active_decks, selected_deck_id=selected_deck_id)

@app.route('/api/deck/<int:deck_id>/cards')
@login_required
def api_deck_cards(deck_id):
    deck = Deck.query.filter_by(id=deck_id, user_id=current_user.id).first()
    if not deck:
        return jsonify({"success": False, "error": "Deck not found"}), 404
        
    cards = Card.query.filter_by(deck_id=deck.id).all()
    cards_data = [{
        "id": c.id,
        "topic": c.topic,
        "question": c.question,
        "answer": c.answer,
        "hint": c.hint,
        "example": c.example,
        "difficulty": c.difficulty
    } for c in cards]
    
    return jsonify({"success": True, "cards": cards_data})

@app.route('/card/<int:card_id>/status', methods=['POST'])
@login_required
def update_card_status(card_id):
    card = Card.query.join(Deck).filter(Card.id == card_id, Deck.user_id == current_user.id).first_or_404()
    
    data = request.get_json() or {}
    status = data.get('status')
    
    if status in ['known', 'review', 'unseen']:
        card.known_status = status
        card.last_reviewed = datetime.utcnow()
        db.session.commit()
        return jsonify({"success": True})
        
    return jsonify({"success": False, "error": "Invalid status"}), 400

@app.route('/quiz/submit', methods=['POST'])
@login_required
def quiz_submit():
    data = request.get_json() or {}
    deck_id = data.get('deck_id')
    score = data.get('score')
    total = data.get('total')
    mode = data.get('mode')
    
    deck = Deck.query.filter_by(id=deck_id, user_id=current_user.id).first()
    if not deck or score is None or not total:
        return jsonify({"success": False, "error": "Invalid request parameters"}), 400
        
    session_log = QuizHistory(
        user_id=current_user.id,
        deck_id=deck.id,
        score=score,
        total_questions=total,
        mode=mode
    )
    db.session.add(session_log)
    db.session.commit()
    
    return jsonify({"success": True})

@app.route('/deck/<int:deck_id>/delete', methods=['POST'])
@login_required
def delete_deck(deck_id):
    deck = Deck.query.filter_by(id=deck_id, user_id=current_user.id).first_or_404()
    db.session.delete(deck)
    db.session.commit()
    flash('Deck deleted successfully.', 'info')
    return redirect(url_for('dashboard'))

# ──────────────────────────────────────────────────────────
# UPLOAD DOCUMENT & CARD PROCESSOR
# ──────────────────────────────────────────────────────────

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "error": "No file selected"}), 400
        
    ext = file.filename.split('.')[-1].lower()
    if ext not in ['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg']:
        return jsonify({"success": False, "error": "File extension not allowed"}), 400
        
    try:
        file_bytes = file.read()
        print(f"[*] Upload file triggered: filename={file.filename}, size={len(file_bytes)} bytes")
        
        # 1. Parse File Content
        if ext == 'pdf':
            segments = parse_pdf(file_bytes)
        elif ext == 'docx':
            segments = parse_docx(file_bytes)
        elif ext in ['png', 'jpg', 'jpeg']:
            segments = parse_image(file_bytes)
        else: # TXT
            full_text = file_bytes.decode('utf-8', errors='ignore')
            segments = chunk_by_length(full_text)

        print(f"[*] Extracted segments count: {len(segments)}")
        if not segments:
            print("[!] No segments extracted from file.")
            return jsonify({"success": False, "error": "Could not extract text from document."}), 400

        print(f"[*] Segment lengths: {[len(s) for s in segments]}")

        # Create new study deck
        deck_name = file.filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
        deck = Deck(user_id=current_user.id, name=deck_name)
        db.session.add(deck)
        db.session.commit()

        # 2. Loop segments and run LLM Generation
        generated_cards_count = 0
        
        # Process first 5 segments max to prevent excessive token delays in demo
        for idx, text_block in enumerate(segments[:5]):
            if len(text_block.strip()) < 40:
                print(f"[*] Skipping segment {idx} due to length: {len(text_block.strip())}")
                continue
                
            print(f"[*] Generating cards for segment {idx} (length={len(text_block)})")
            cards = generate_cards_ai(text_block, deck_name)
            print(f"[*] generate_cards_ai returned {len(cards)} cards for segment {idx}")
            for c_data in cards:
                c_topic = (c_data.get('topic') or 'General').strip()[:100]
                c_question = (c_data.get('question') or '').strip()
                c_answer = (c_data.get('answer') or '').strip()
                c_hint = (c_data.get('hint') or '').strip() or None
                c_example = (c_data.get('example') or '').strip() or None
                c_difficulty = (c_data.get('difficulty') or 'Medium').strip()

                card = Card(
                    deck_id=deck.id,
                    topic=c_topic,
                    question=c_question,
                    answer=c_answer,
                    hint=c_hint,
                    example=c_example,
                    difficulty=c_difficulty
                )
                
                # Check formatting validation
                if card.question and card.answer:
                    if card.difficulty not in ['Easy', 'Medium', 'Hard']:
                        card.difficulty = 'Medium'
                    db.session.add(card)
                    generated_cards_count += 1
                else:
                    print(f"[!] Invalid card skipped: question='{card.question}', answer='{card.answer}'")
                    
        print(f"[*] Finished processing segments. Total valid cards generated: {generated_cards_count}")
        if generated_cards_count == 0:
            print("[!] No cards were generated. Deleting deck and returning 400.")
            db.session.delete(deck)
            db.session.commit()
            return jsonify({"success": False, "error": "AI could not generate valid flashcards. Verify document text content."}), 400
            
        db.session.commit()
        print(f"[*] Successfully created deck {deck.name} (id={deck.id}) with {generated_cards_count} cards.")
        return jsonify({"success": True, "deck_id": deck.id})
        
    except Exception as e:
        import traceback
        with open("upload_crash.txt", "w") as f:
            f.write(traceback.format_exc())
        print(f"[!] Upload processing crash: {e}")
        print(traceback.format_exc())
        return jsonify({"success": False, "error": f"Internal process error: {str(e)}"}), 500

# ──────────────────────────────────────────────────────────
# REPORTLAB PDF GENERATOR EXPORT
# ──────────────────────────────────────────────────────────

@app.route('/export/pdf/<int:deck_id>')
@login_required
def export_pdf(deck_id):
    deck = Deck.query.filter_by(id=deck_id, user_id=current_user.id).first_or_404()
    cards = Card.query.filter_by(deck_id=deck.id).all()
    
    # Create PDF in-memory buffer
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        rightMargin=40, leftMargin=40, topMargin=45, bottomMargin=45
    )
    
    styles = getSampleStyleSheet()
    
    # Custom color themes
    navy = colors.HexColor('#0d1526')
    cyan = colors.HexColor('#00d4ff')
    grey = colors.HexColor('#64748b')
    light_bg = colors.HexColor('#f8fafc')
    
    title_style = ParagraphStyle(
        'DeckTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        textColor=navy,
        spaceAfter=15
    )
    
    meta_style = ParagraphStyle(
        'DeckMeta',
        fontName='Helvetica',
        fontSize=10,
        textColor=grey,
        spaceAfter=25
    )
    
    q_title_style = ParagraphStyle(
        'CardQTitle',
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=cyan,
        spaceAfter=6
    )
    
    q_text_style = ParagraphStyle(
        'CardQText',
        fontName='Helvetica',
        fontSize=10.5,
        textColor=navy,
        leading=14
    )
    
    a_title_style = ParagraphStyle(
        'CardATitle',
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=colors.HexColor('#a855f7'),
        spaceAfter=6
    )
    
    a_text_style = ParagraphStyle(
        'CardAText',
        fontName='Helvetica',
        fontSize=10.5,
        textColor=navy,
        leading=14
    )
    
    story = []
    
    # Document header
    story.append(Paragraph(f"IntelliCard AI — Study Deck", title_style))
    story.append(Paragraph(f"Deck Name: {deck.name} | Total Cards: {len(cards)} | Exported: {datetime.now().strftime('%Y-%m-%d')}", meta_style))
    
    # Build a table list of cards
    table_data = []
    for idx, c in enumerate(cards):
        # Front Column: Topic, Difficulty, Question, Hint
        q_block = [
            Paragraph(f"CARD #{idx + 1} — {c.topic.upper()} [{c.difficulty}]", q_title_style),
            Spacer(1, 4),
            Paragraph(c.question, q_text_style),
        ]
        if c.hint:
            hint_style = ParagraphStyle('HintStyle', fontName='Helvetica-Oblique', fontSize=9, textColor=grey)
            q_block.append(Spacer(1, 6))
            q_block.append(Paragraph(f"Hint: {c.hint}", hint_style))
            
        # Back Column: Answer & Examples
        a_block = [
            Paragraph("EXPLANATION / ANSWER", a_title_style),
            Spacer(1, 4),
            Paragraph(c.answer, a_text_style),
        ]
        if c.example:
            ex_style = ParagraphStyle('ExStyle', fontName='Courier', fontSize=8.5, textColor=colors.HexColor('#1e293b'))
            a_block.append(Spacer(1, 6))
            a_block.append(Paragraph(c.example.replace('\n', '<br/>'), ex_style))
            
        table_data.append([q_block, a_block])
        
    if table_data:
        # Create Table. 2 columns, width 260 points each (total width = 520 points)
        card_table = Table(table_data, colWidths=[260, 260])
        card_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('PADDING', (0, 0), (-1, -1), 16),
            ('BACKGROUND', (0, 0), (0, -1), light_bg),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
            ('TOPPADDING', (0, 0), (-1, -1), 14),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 14),
        ]))
        story.append(card_table)
    else:
        story.append(Paragraph("This deck contains no flashcards.", q_text_style))
        
    doc.build(story)
    pdf_buffer.seek(0)
    
    clean_filename = re.sub(r'[^a-zA-Z0-9]', '_', deck.name.lower())
    return send_file(
        pdf_buffer,
        as_attachment=True,
        download_name=f"intellicard_{clean_filename}.pdf",
        mimetype="application/pdf"
    )

# ──────────────────────────────────────────────────────────
# INITIALIZE SERVER & DB TABLES
# ──────────────────────────────────────────────────────────

with app.app_context():
    db.create_all()

if __name__ == '__main__':
    print("--------------------------------------------------")
    print(" IntelliCard AI server database created successfully.")
    print(" Starting development server at http://127.0.0.1:5000")
    print("--------------------------------------------------")
    
    app.run(debug=True, port=5000)
